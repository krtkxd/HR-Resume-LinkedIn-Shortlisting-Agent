"""
utils/resume_parser.py
──────────────────────
Extracts raw text from PDF and DOCX resume files.

Supports:
  - PDF via pdfplumber (primary) with PyMuPDF fallback
  - DOCX via python-docx
  - Plain text (.txt)

Returns a clean, whitespace-normalised string.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Union


# ──────────────────────────────────────────────
# Public interface
# ──────────────────────────────────────────────

def extract_text(source: Union[str, Path, bytes], filename: str = "") -> str:
    """
    Extract raw text from a resume file.

    Args:
        source   : file path (str/Path) OR raw bytes
        filename : original filename (used to determine type when bytes passed)

    Returns:
        Cleaned plain-text string
    """
    if isinstance(source, (str, Path)):
        path = Path(source)
        suffix = path.suffix.lower()
        raw_bytes = path.read_bytes()
    else:
        raw_bytes = source
        suffix = Path(filename).suffix.lower() if filename else ""

    if suffix == ".pdf":
        text = _extract_pdf(raw_bytes)
    elif suffix in (".docx", ".doc"):
        text = _extract_docx(raw_bytes)
    elif suffix == ".txt":
        text = raw_bytes.decode("utf-8", errors="replace")
    else:
        # Try PDF then DOCX as fallback
        try:
            text = _extract_pdf(raw_bytes)
        except Exception:
            try:
                text = _extract_docx(raw_bytes)
            except Exception:
                text = raw_bytes.decode("utf-8", errors="replace")

    return _clean_text(text)


# ──────────────────────────────────────────────
# PDF extraction
# ──────────────────────────────────────────────

def _extract_pdf(raw_bytes: bytes) -> str:
    """Try pdfplumber first; fall back to PyMuPDF."""
    try:
        return _pdf_pdfplumber(raw_bytes)
    except Exception:
        pass
    try:
        return _pdf_pymupdf(raw_bytes)
    except Exception:
        raise RuntimeError("Could not extract text from PDF. "
                           "Ensure pdfplumber or PyMuPDF is installed.")


def _pdf_pdfplumber(raw_bytes: bytes) -> str:
    import pdfplumber
    pages = []
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n".join(pages)


def _pdf_pymupdf(raw_bytes: bytes) -> str:
    import fitz  # PyMuPDF
    pages = []
    doc = fitz.open(stream=raw_bytes, filetype="pdf")
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return "\n".join(pages)


# ──────────────────────────────────────────────
# DOCX extraction
# ──────────────────────────────────────────────

def _extract_docx(raw_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(raw_bytes))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


# ──────────────────────────────────────────────
# Text cleaning
# ──────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Remove excessive whitespace and non-printable characters."""
    # Replace form-feeds, vertical tabs, etc. with newline
    text = re.sub(r"[\f\v\r]+", "\n", text)
    # Collapse multiple blank lines to a maximum of two
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove non-printable characters except newlines and tabs
    text = re.sub(r"[^\x09\x0a\x20-\x7e\u00a0-\uffff]", " ", text)
    # Collapse horizontal whitespace runs
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()
